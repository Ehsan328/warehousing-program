from docx import Document

def history_table_maker(fo):
    doc= Document('C:\\Ehsan\\history.docx')
    
    namemalek= fo[0]
    namekala= fo[1]
    vazn= fo[2]
    tedad= fo[3]
    vk= fo[4]
    tarikh= fo[5]

    table_header= [namemalek, 'نام کالا', 'وزن', 'تعداد', 'ورود/خوروج', 'تاریخ']

    bar= [
        ['',namekala , vazn, tedad, vk, tarikh],
    ]

    table= doc.add_table(rows= 1, cols= 6)

    for i in range(6):
        table.rows[0].cells[i].text= table_header[i]

    for namemalek, namekala, vazn, tedad, vk, tarikh in bar:
        cells= table.add_row().cells
        cells[0].text= '1'
        cells[1].text= namekala
        cells[2].text= str(vazn)
        cells[3].text= str(tedad)
        cells[4].text= str(vk)
        cells[5].text= str(tarikh)
    doc.add_paragraph('--------------------------------------------------------------------------------------------------------------------------------------').bold= True

    doc.save('C:\\Ehsan\\history.docx')
