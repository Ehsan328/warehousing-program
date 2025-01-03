from docx import Document

def tablemake(malek, vorodkhoroj):
    doc= Document('C:\\Ehsan\\inventory.docx')
    b= []
    for i in malek.keys():
        b.append(i)
    
    namemalek= b[0]
    namekala= b[1]
    vazn= malek[b[1]][0]
    tedad= malek[b[1]][1]

    table_header= [namemalek, 'نام کالا'+ ' .'+ 'یکا', 'وزن', 'تعداد']

    bar= [
        ['',namekala , vazn, tedad],
    ]

    table= doc.add_table(rows= 1, cols= 4)

    for i in range(4):
        table.rows[0].cells[i].text= table_header[i]
        
    if vorodkhoroj== 'vorod':
        for namemalek, namekala, vazn, tedad in bar:
            cells= table.add_row().cells
            cells[0].text= ''
            cells[1].text= namekala
            cells[2].text= str(vazn)
            cells[3].text= str(tedad)
        doc.add_paragraph('--------------------------------------------------------------------------------------------------------------------------------------').bold= True
    if vorodkhoroj== 'khoroj':
        for namemalek, namekala, vazn, tedad in bar:
            cells= table.add_row().cells
            cells[0].text= ''
            cells[1].text= namekala
            cells[2].text= str(-int(vazn))
            cells[3].text= str(-int(tedad))
        doc.add_paragraph('--------------------------------------------------------------------------------------------------------------------------------------').bold= True


        
    doc.save('C:\\Ehsan\\inventory.docx')
    return None

