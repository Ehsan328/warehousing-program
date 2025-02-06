from docx import Document
from history_table_maker import history_table_maker

def history (vorodkhoroj, tarikh, namemalek, namekala, vazn, tedad, yeka):
    doc= Document('C:\\Ehsan\\history.docx')
    
    if vorodkhoroj== 'vorod':
        vorodkhoroj= 'ورودی'
    if vorodkhoroj== 'khoroj':
        vorodkhoroj= 'خروجی'
    
    namekala+=' .' + yeka
    
    # fo is a list of information for register
    fo= [namemalek, namekala, vazn, tedad, vorodkhoroj, tarikh]
    
    #creat a list of owners
    listmalek= []
    for table in doc.tables:
        listmalek.append(table.cell(0, 0).text)
    #if owner isn't exist make a table for him
    if namemalek not in listmalek:
        return history_table_maker(fo)
    
    #if owner isn't exist make a row
    else:
        for table in doc.tables:
            if table.cell(0, 0).text== namemalek:
                im= [[str(len(table.rows)-1), namekala, vazn, tedad, vorodkhoroj, tarikh]]
                for namemalek, kala, vazn, tedad, vorodkhoroj, tarikh in im:
                    cells= table.add_row().cells
                    cells[0].text= str(len(table.rows)-1)
                    cells[1].text= str(namekala)
                    cells[2].text= str(vazn)
                    cells[3].text= str(tedad)
                    cells[4].text= str(vorodkhoroj)
                    cells[5].text= str(tarikh)

    
    doc.save('C:\\Ehsan\\history.docx')


