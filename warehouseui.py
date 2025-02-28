# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'warehouseui.ui'
#
# Created by: PyQt5 UI code generator 5.15.11
#
# WARNING: Any manual changes made to this file will be lost when pyuic5 is
# run again.  Do not edit this file unless you know what you are doing.


from PyQt5 import QtCore, QtGui, QtWidgets
from warehouse import warehouse
from history import history
from error import error
from docx import Document
from time import sleep

import sys
from PyQt5.QtWidgets import QDialog, QApplication
from PyQt5.uic import loadUi



class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(829, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.tabWidget = QtWidgets.QTabWidget(self.centralwidget)
        self.tabWidget.setGeometry(QtCore.QRect(0, 0, 821, 551))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.tabWidget.setFont(font)
        self.tabWidget.setCursor(QtGui.QCursor(QtCore.Qt.ArrowCursor))
        self.tabWidget.setAutoFillBackground(False)
        self.tabWidget.setUsesScrollButtons(True)
        self.tabWidget.setDocumentMode(True)
        self.tabWidget.setObjectName("tabWidget")
        self.jadid = QtWidgets.QWidget()
        self.jadid.setObjectName("jadid")
        self.vaznline = QtWidgets.QLineEdit(self.jadid)
        self.vaznline.setGeometry(QtCore.QRect(360, 210, 121, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.vaznline.setFont(font)
        self.vaznline.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.vaznline.setObjectName("vaznline")
        self.namemalekline = QtWidgets.QLineEdit(self.jadid)
        self.namemalekline.setGeometry(QtCore.QRect(640, 160, 171, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.namemalekline.setFont(font)
        self.namemalekline.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.namemalekline.setText("")
        self.namemalekline.setObjectName("namemalekline")
        self.yekaline = QtWidgets.QLineEdit(self.jadid)
        self.yekaline.setGeometry(QtCore.QRect(310, 240, 51, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.yekaline.setFont(font)
        self.yekaline.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.yekaline.setObjectName("yekaline")
        self.namekalaline = QtWidgets.QLineEdit(self.jadid)
        self.namekalaline.setGeometry(QtCore.QRect(50, 160, 171, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.namekalaline.setFont(font)
        self.namekalaline.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.namekalaline.setObjectName("namekalaline")
        self.khorojradio = QtWidgets.QRadioButton(self.jadid)
        self.khorojradio.setGeometry(QtCore.QRect(750, 70, 61, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.khorojradio.setFont(font)
        self.khorojradio.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.khorojradio.setObjectName("khorojradio")
        self.vorodradio = QtWidgets.QRadioButton(self.jadid)
        self.vorodradio.setGeometry(QtCore.QRect(750, 40, 51, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        font.setBold(False)
        font.setWeight(50)
        self.vorodradio.setFont(font)
        self.vorodradio.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.vorodradio.setCheckable(True)
        self.vorodradio.setChecked(False)
        self.vorodradio.setObjectName("vorodradio")
        self.tedadline = QtWidgets.QLineEdit(self.jadid)
        self.tedadline.setGeometry(QtCore.QRect(360, 240, 121, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.tedadline.setFont(font)
        self.tedadline.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.tedadline.setObjectName("tedadline")
        self.tarikhline = QtWidgets.QLineEdit(self.jadid)
        self.tarikhline.setGeometry(QtCore.QRect(520, 50, 161, 41))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.tarikhline.setFont(font)
        self.tarikhline.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.tarikhline.setText("")
        self.tarikhline.setFrame(True)
        self.tarikhline.setObjectName("tarikhline")
        self.errorlabel = QtWidgets.QLabel(self.jadid)
        self.errorlabel.setGeometry(QtCore.QRect(180, 330, 451, 31))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.errorlabel.setFont(font)
        self.errorlabel.setStyleSheet("color: rgb(197, 0, 0);")
        self.errorlabel.setText("")
        self.errorlabel.setObjectName("errorlabel")
        self.sabtbut = QtWidgets.QPushButton(self.jadid)
        self.sabtbut.setGeometry(QtCore.QRect(365, 390, 90, 41))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.sabtbut.setFont(font)
        self.sabtbut.setAutoDefault(False)
        self.sabtbut.setDefault(False)
        self.sabtbut.setFlat(False)
        self.sabtbut.setObjectName("sabtbut")
        self.tabWidget.addTab(self.jadid, "")
        self.info = QtWidgets.QWidget()
        self.info.setObjectName("info")
        self.tarikhchebut = QtWidgets.QPushButton(self.info)
        self.tarikhchebut.setGeometry(QtCore.QRect(200, 10, 92, 31))
        self.tarikhchebut.setObjectName("tarikhchebut")
        self.namemalekcombo = QtWidgets.QComboBox(self.info)
        self.namemalekcombo.setGeometry(QtCore.QRect(10, 30, 171, 31))
        self.namemalekcombo.setObjectName("namemalekcombo")
        self.pushButton = QtWidgets.QPushButton(self.info)
        self.pushButton.setGeometry(QtCore.QRect(200, 56, 92, 31))
        self.pushButton.setObjectName("pushButton")
        self.tableWidget = QtWidgets.QTableWidget(self.info)
        self.tableWidget.setGeometry(QtCore.QRect(0, 91, 821, 431))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.tableWidget.setFont(font)
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setColumnCount(5)
        self.tableWidget.setRowCount(0)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(3, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(4, item)
        self.tableWidget.horizontalHeader().setDefaultSectionSize(160)
        self.hazf = QtWidgets.QPushButton(self.info)
        self.hazf.setGeometry(QtCore.QRect(320, 30, 92, 31))
        self.hazf.setObjectName("hazf")
        self.tabWidget.addTab(self.info, "")
        self.tab = QtWidgets.QWidget()
        self.tab.setObjectName("tab")
        self.label = QtWidgets.QLabel(self.tab)
        self.label.setGeometry(QtCore.QRect(10, 10, 511, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.tab)
        self.label_2.setGeometry(QtCore.QRect(10, 49, 501, 41))
        self.label_2.setObjectName("label_2")
        self.tabWidget.addTab(self.tab, "")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 829, 26))
        self.menubar.setObjectName("menubar")
        self.menuEhsan = QtWidgets.QMenu(self.menubar)
        self.menuEhsan.setObjectName("menuEhsan")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.menubar.addAction(self.menuEhsan.menuAction())

        self.retranslateUi(MainWindow)
        self.tabWidget.setCurrentIndex(0)
        
        ### put the buttons here ###
        
        self.sabtbut.clicked.connect(self.sabtjadid) # type: ignore
        self.pushButton.clicked.connect(self.mojodi)
        self.tarikhchebut.clicked.connect(self.tarikhche)
        self.hazf.clicked.connect(self.hazfmalek)

        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        MainWindow.setTabOrder(self.vorodradio, self.khorojradio)
        MainWindow.setTabOrder(self.khorojradio, self.tarikhline)
        MainWindow.setTabOrder(self.tarikhline, self.namemalekline)
        MainWindow.setTabOrder(self.namemalekline, self.namekalaline)
        MainWindow.setTabOrder(self.namekalaline, self.vaznline)
        MainWindow.setTabOrder(self.vaznline, self.tedadline)
        MainWindow.setTabOrder(self.tedadline, self.yekaline)
        MainWindow.setTabOrder(self.yekaline, self.sabtbut)
        MainWindow.setTabOrder(self.sabtbut, self.tabWidget)
        MainWindow.setTabOrder(self.tabWidget, self.namemalekcombo)
        MainWindow.setTabOrder(self.namemalekcombo, self.tarikhchebut)
        MainWindow.setTabOrder(self.tarikhchebut, self.pushButton)
        
        #give combo name of owners
        doc = Document('C:\\Ehsan\\inventory.docx')
        listnamemalek= []
        for table in doc.tables:
            listnamemalek.append(table.cell(0,0).text)
            
        self.namemalekcombo.insertItems(1, listnamemalek)
    
    #remove the name
    def hazfmalek(self):
        historydoc = Document('C:\\Ehsan\\history.docx')
        inventorydoc = Document('C:\\Ehsan\\inventory.docx')
        name = self.namemalekcombo.currentText()
        index = self.namemalekcombo.findText(name)
        
        for i in historydoc.tables:
            if i.cell(0,0).text == name:
                i._element.getparent().remove(i._element)
                
        for i in inventorydoc.tables:
            if i.cell(0,0).text == name:
                i._element.getparent().remove(i._element)
        
        self.namemalekcombo.removeItem(index)
        
        historydoc.save('C:\\Ehsan\\history.docx')
        inventorydoc.save('C:\\Ehsan\\inventory.docx')

        
        
    
    
        
    #when you press the sabt button you'll connect to this function
    def sabtjadid(self):
        
        #get the list of owners from docx
        doc = Document('C:\\Ehsan\\history.docx')
        listnamemalek= []
        for table in doc.tables:
            listnamemalek.append(table.cell(0,0).text)
        
        #if name of owner isn't in the list owner mention it
        vazeiatmalek= ''
        if self.namemalekline.text() not in listnamemalek:
            vazeiatmalek= '(نام مالک جدید است)'



        
        #errors to fill the lines correctly
        self.errorlabel.setText('')
        if not(self.vorodradio.isChecked() or self.khorojradio.isChecked()):
            self.errorlabel.setText('ورودی یا خروجی بودن بار را تایین کنید')
            return ''
        elif self.tarikhline.text()== '':
            self.errorlabel.setText('کادر تاریخ را پر کنید')
            return ''
        elif not(error('tarikh',self.tarikhline.text())):
            self.errorlabel.setText('فرمت تاریخ اشتباه است')
            return ''
        elif self.namemalekline.text()== '':
            self.errorlabel.setText('کادر نام مالک را پر کنید')
            return ''
        elif self.namekalaline.text()== '':
            self.errorlabel.setText('کادر نام کالا را پر کنید')
            return ''
        elif self.vaznline.text()== '':
            self.errorlabel.setText('کادر وزن را پر کنید')
            return ''
        elif not(error('vazn',self.vaznline.text())):
            self.errorlabel.setText('فرمت وزن اشتباه است')
            return ''
        elif self.tedadline.text()== '':
            self.errorlabel.setText('کادر تعداد را پر کنید')
            return ''
        elif not(error('tedad',self.tedadline.text())):
            self.errorlabel.setText('فرمت تعداد اشتباه است')
            return ''
        elif self.yekaline.text()== '':
            self.errorlabel.setText('کادر یکا را پر کنید')
            return ''
        
        #put the informatioin in the docxes        
        if self.vorodradio.isChecked():
            vorodkhoroj= 'vorod'
        if self.khorojradio.isChecked():
            vorodkhoroj= 'khoroj'
        warehouse(vorodkhoroj, self.tarikhline.text(), self.namemalekline.text(), self.namekalaline.text()
                , self.vaznline.text(), self.tedadline.text(), self.yekaline.text()
)
        history(vorodkhoroj, self.tarikhline.text(), self.namemalekline.text(), self.namekalaline.text()
                , self.vaznline.text(), self.tedadline.text(), self.yekaline.text()
)
        self.namemalekcombo.clear()
        self.namemalekcombo.insertItems(1, listnamemalek)
        
        sleep(0.5)
        self.tarikhline.setText('')
        self.namemalekline.setText('')
        self.namekalaline.setText('')
        self.vaznline.setText('')
        self.tedadline.setText('')
        self.yekaline.setText('')
        
        #get the list of owners from docx
        doc = Document('C:\\Ehsan\\history.docx')
        listnamemalek= []
        for table in doc.tables:
            listnamemalek.append(table.cell(0,0).text)
        
        self.namemalekcombo.clear()
        self.namemalekcombo.insertItems(1, listnamemalek)

        
        

                
        self.errorlabel.setText('اطلاعات با موفقیت ثبت شد'+ ' '+ vazeiatmalek)
    
    def mojodi(self):
        doc = Document('C:\\Ehsan\\inventory.docx')
        name= self.namemalekcombo.currentText()
        for table in doc.tables:
            if table.cell(0,0).text== name:
                self.tableWidget.setRowCount(len(table.rows)-1)
                for i in range(len(table.rows)-1):
                    self.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(table.cell(i+1, 1).text))    
                    self.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(table.cell(i+1, 2).text))
                    self.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(table.cell(i+1, 3).text))
                    self.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem('---'))
                    self.tableWidget.setItem(i, 4, QtWidgets.QTableWidgetItem('---'))
    
    def tarikhche(self):
        doc = Document('C:\\Ehsan\\history.docx')
        name= self.namemalekcombo.currentText()
        for table in doc.tables:
            if table.cell(0,0).text== name:
                self.tableWidget.setRowCount(len(table.rows)-1)
                for i in range(len(table.rows)-1):
                    self.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(table.cell(i+1, 1).text))    
                    self.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(table.cell(i+1, 2).text))
                    self.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(table.cell(i+1, 3).text))
                    self.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem(table.cell(i+1, 4).text))
                    self.tableWidget.setItem(i, 4, QtWidgets.QTableWidgetItem(table.cell(i+1, 5).text))


    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.vaznline.setPlaceholderText(_translate("MainWindow", "وزن (kg)"))
        self.namemalekline.setPlaceholderText(_translate("MainWindow", "نام مالک"))
        self.yekaline.setPlaceholderText(_translate("MainWindow", "یکا"))
        self.namekalaline.setPlaceholderText(_translate("MainWindow", "نام کالا"))
        self.khorojradio.setText(_translate("MainWindow", "خروج"))
        self.vorodradio.setText(_translate("MainWindow", "ورود"))
        self.tedadline.setPlaceholderText(_translate("MainWindow", "تعداد"))
        self.tarikhline.setPlaceholderText(_translate("MainWindow", "تاریخ (1403/09/12)"))
        self.sabtbut.setText(_translate("MainWindow", "ثبت"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.jadid), _translate("MainWindow", "جدید"))
        self.tarikhchebut.setText(_translate("MainWindow", "تاریخچه"))
        self.pushButton.setText(_translate("MainWindow", "موجودی"))
        item = self.tableWidget.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "نام کالا. یکا"))
        item = self.tableWidget.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "وزن"))
        item = self.tableWidget.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "تعداد"))
        item = self.tableWidget.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "ورود/خروج"))
        item = self.tableWidget.horizontalHeaderItem(4)
        item.setText(_translate("MainWindow", "تاریخ"))
        self.hazf.setText(_translate("MainWindow", "حذف مالک"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.info), _translate("MainWindow", "اطلاعات"))
        self.label.setText(_translate("MainWindow", "برای سفارش پروژه های برنامه نویسی بیشتر به این ادرس ایمیل پیام بدهید."))
        self.label_2.setText(_translate("MainWindow", "programmingorder328@gmail.com"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab), _translate("MainWindow", "درباره ما"))
        self.menuEhsan.setTitle(_translate("MainWindow", "Ehsan"))


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
