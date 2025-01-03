from docx import Document
from tablemaker import tablemake

#کالا های هم نام با یکای متفاوت از هم متمایز می شن با فرمت کالا+ +یکا


def warehouse (vorodkhoroj, tarikh, namemalek, namekala, vazn, tedad, yeka):
    #openning inventory word file
    doc = Document('C:\\Ehsan\\inventory.docx')
    kala= namekala+ ' .'+ yeka

    malek= {namemalek: '', kala: [vazn, tedad]}
    b= []
    for i in malek.keys():
        b.append(i)

    #creat a list of name of goods
    listkala= []
    for table in doc.tables:
        if table.cell(0, 0).text== namemalek:
            for i in range(1, len(table.rows)):
                listkala.append(table.cell(i, 1).text)

    #creat a list of owners
    listmalek= []
    for table in doc.tables:
        listmalek.append(table.cell(0, 0).text)

    #if owner not exist make a table for him    
    if namemalek not in listmalek:
        return tablemake (malek, vorodkhoroj)

    else:

        for table in doc.tables:
            if table.cell(0, 0).text== namemalek:

                #if goods not exist make a row for it
                if kala not in listkala:
                    im= [['', kala, vazn, tedad]]
                    if vorodkhoroj== 'vorod':                    
                        for namemalek, kala, vazn, tedad in im:
                            cells= table.add_row().cells
                            cells[0].text= ''
                            cells[1].text= kala
                            cells[2].text= str(vazn)
                            cells[3].text= str(tedad)
                    if vorodkhoroj== 'khoroj':
                        for namemalek, kala, vazn, tedad in im:
                            cells= table.add_row().cells
                            cells[0].text= ''
                            cells[1].text= kala
                            cells[2].text= str(-int(vazn))
                            cells[3].text= str(-int(tedad))
                    
                                
                #if owner is exist and goods is exist too 
                # so increase or decrease inventory considering entry or exit 
                elif kala in listkala:
                    #entry
                    if vorodkhoroj== 'vorod':
                        for i in range(len(table.rows)):
                            if table.cell(i, 1).text== kala:
                                table.cell(i, 2).text= str(int(table.cell(i, 2).text)+ int(vazn))
                                table.cell(i, 3).text= str(int(table.cell(i, 3).text)+ int(tedad))
                    # quit
                    if vorodkhoroj== 'khoroj':
                        for i in range(len(table.rows)):
                            if table.cell(i, 1).text== kala:
                                table.cell(i, 2).text= str(int(table.cell(i, 2).text)- int(vazn))
                                table.cell(i, 3).text= str(int(table.cell(i, 3).text)- int(tedad))
                            
                        
    #saving inventory word file
    doc.save('C:\\Ehsan\\inventory.docx')

#namemalek = "علی موئمنی"
#tarikh = None
#namekala = 'مرغ'
#vazn = 33
#tedad = 3
#yeka = 'کارتن'
#vorodkhoroj = 'khoroj'
#warehouse(vorodkhoroj, tarikh, namemalek, namekala, vazn, tedad, yeka)
